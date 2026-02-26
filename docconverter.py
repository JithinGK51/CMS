import os
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_margins(cell, top=0, start=0, bottom=0, end=0):
    """Sets the margins for a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m in [('top', top), ('start', start), ('bottom', bottom), ('end', end)]:
        if m[1] is not None:
            node = OxmlElement(f'w:{m[0]}')
            node.set(qn('w:w'), str(m[1]))
            node.set(qn('w:type'), 'dxa')
            tcMar.append(node)
    tcPr.append(tcMar)

def create_professional_doc():
    doc = Document()
    
    # --- Style Configuration ---
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # --- COVER PAGE ---
    doc.add_paragraph("\n" * 5)
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("PUBLIC COMPLAINT MANAGEMENT SYSTEM")
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = RGBColor(41, 128, 185)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Comprehensive Technical Design & System Documentation")
    run.font.size = Pt(16)
    run.font.italic = True

    doc.add_paragraph("\n" * 10)
    
    info_table = doc.add_table(rows=4, cols=2)
    info_table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    data = [
        ["Project Version:", "v1.0.4"],
        ["Organization:", "Public Works Authority / Internal Systems Div."],
        ["Author:", "Senior Systems Architect - Antigravity Systems"],
        ["Date:", datetime.now().strftime("%B %d, %Y")]
    ]
    
    for i, row in enumerate(data):
        info_table.cell(i, 0).text = row[0]
        info_table.cell(i, 1).text = row[1]
        info_table.cell(i, 0).paragraphs[0].runs[0].font.bold = True

    doc.add_page_break()

    # --- TABLE OF CONTENTS ---
    doc.add_heading('Table of Contents', level=1)
    doc.add_paragraph("1. Abstract\n2. Introduction\n3. Existing vs. Proposed System\n4. System Objectives\n5. Technical Architecture\n6. Module Specifications\n7. Database Architecture\n8. System Workflow & Logic\n9. User Roles & Permissions\n10. Deployment Requirements\n11. System Interfaces\n12. Future Scope & Conclusion")
    doc.add_page_break()

    # --- 1. ABSTRACT ---
    doc.add_heading('1. Abstract', level=1)
    abstract = (
        "The Public Complaint Management System (PCMS) is a mission-critical digital infrastructure "
        "designed to bridge the communication gap between citizens and public service providers. "
        "By leveraging a serverless architecture with Flask and Supabase, the system provides a "
        "high-availability, transparent, and secure environment for registering, tracking, and resolving "
        "public grievances. Key features include tokenized tracking, role-based access control (RBAC), "
        "and multi-media evidence handling."
    )
    doc.add_paragraph(abstract)

    # --- 2. INTRODUCTION ---
    doc.add_heading('2. Introduction', level=1)
    intro = (
        "In the modern era of governance, efficiency and transparency are paramount. Traditional "
        "paper-based complaint systems are often marred by delays, lack of accountability, and poor "
        "data visibility. PCMS centralizes the workflow into a streamlined digital pipeline, ensuring "
        "that every voice is heard and every complaint is actionable."
    )
    doc.add_paragraph(intro)

    # --- 3. EXISTING VS PROPOSED ---
    doc.add_heading('3. Existing System Review', level=1)
    doc.add_paragraph("Analysis of the legacy infrastructure reveals significant bottlenecks:")
    legacy = doc.add_paragraph()
    legacy.add_run("• Manual Data Entry: ").bold = True
    legacy.add_run("High error rates and labor-intensive processes.\n")
    legacy.add_run("• Opaque Tracking: ").bold = True
    legacy.add_run("Citizens have no visibility into the resolution status.\n")
    legacy.add_run("• Decentralized Records: ").bold = True
    legacy.add_run("Information silos preventing effective resource allocation.")

    doc.add_heading('4. Proposed Solution Architecture', level=1)
    proposed = doc.add_paragraph()
    proposed.add_run("• Real-time Tracking: ").bold = True
    proposed.add_run("Token-based system for status transparency.\n")
    proposed.add_run("• Automated Routing: ").bold = True
    proposed.add_run("Smart assignment based on departmental categories.\n")
    proposed.add_run("• Centralized PostgreSQL: ").bold = True
    proposed.add_run("Unified data source for analytics and reporting.")

    # --- 5. SYSTEM OBJECTIVES ---
    doc.add_heading('5. System Objectives', level=1)
    obj = doc.add_table(rows=1, cols=2)
    obj.cell(0, 0).text = "Primary Objectives"
    obj.cell(0, 1).text = "Secondary Objectives"
    
    p_obj = (
        "• Digitize complaint registration\n"
        "• Implement real-time tracking\n"
        "• Enforce RBAC security\n"
        "• Provide analytical dashboards"
    )
    s_obj = (
        "• Reduce paper footprints\n"
        "• Enhance public trust\n"
        "• Optimize staff workload\n"
        "• Archive historical data"
    )
    obj.cell(0, 0).add_paragraph(p_obj)
    obj.cell(0, 1).add_paragraph(s_obj)

    # --- 6. SYSTEM ARCHITECTURE ---
    doc.add_heading('6. Technical Stack & Architecture', level=1)
    stack = doc.add_table(rows=5, cols=2)
    stack_data = [
        ["Layer", "Technology"],
        ["Presentation Layer", "HTML5, CSS3, Modern JavaScript (Vanilla)"],
        ["Application Tier", "Python 3.x, Flask Web Framework"],
        ["Database Tier", "Supabase PostgreSQL (Cloud-native)"],
        ["Security / Auth", "Supabase Auth (JWT & Role Protection)"]
    ]
    for i, row in enumerate(stack_data):
        stack.cell(i, 0).text = row[0]
        stack.cell(i, 1).text = row[1]
        stack.cell(i, 0).paragraphs[0].runs[0].font.bold = True

    # --- 7. DATABASE ARCHITECTURE ---
    doc.add_heading('7. Database Schema Design', level=1)
    doc.add_paragraph("The system utilizes a relational schema optimized for complaint lifecycle management.")
    
    db_table = doc.add_table(rows=1, cols=3)
    hdr_cells = db_table.rows[0].cells
    hdr_cells[0].text = 'Table Name'
    hdr_cells[1].text = 'Description'
    hdr_cells[2].text = 'Primary Keys / Relations'
    
    tables_info = [
        ["departments", "Organizational units (Police, Water, Electricity, etc.)", "id (UUID)"],
        ["staff", "User accounts with roles (Admin/Staff)", "id (UUID) -> departments"],
        ["categories", "Complaint types (e.g., Street Light, Pothole)", "id (UUID) -> departments"],
        ["complaints", "Core complaint data, citizen info, and status", "id (BigInt), token (Unique)"],
        ["complaint_timeline", "Audit trail for status transitions", "id (UUID) -> complaints"],
        ["complaint_files", "Evidence storage and proof of resolution", "id (UUID) -> complaints"],
        ["notifications", "System alerts for staff assignments", "id (UUID) -> staff"],
        ["system_settings", "Global configurations (prefixes, automation)", "id (PK: 1)"]
    ]
    
    for name, desc, rel in tables_info:
        row_cells = db_table.add_row().cells
        row_cells[0].text = name
        row_cells[1].text = desc
        row_cells[2].text = rel

    # --- 8. PROJECT FLOW DIAGRAM ---
    doc.add_heading('8. System Workflow & Process Logic', level=1)
    flow = doc.add_paragraph()
    flow.add_run("Phase I: Complaint Submission (Public User)").bold = True
    flow.add_run("\nCitizen logs onto portal -> Selects Category -> Fills Description -> Uploads Images -> Submits. System generates unique Token (CMP-YYYY-XXXXX) and notifies Department Admin.")
    
    flow.add_run("\n\nPhase II: Internal Review (Admin)").bold = True
    flow.add_run("\nAdmin reviews incoming queue -> Validates authenticity -> Assigns to specific Staff member. Status updates to 'ASSIGNED'.")
    
    flow.add_run("\n\nPhase III: Resolution (Staff)").bold = True
    flow.add_run("\nStaff receives notification -> Investigates -> Performs task -> Uploads proof of work -> Updates status to 'RESOLVED' or 'IN_PROGRESS'.")
    
    flow.add_run("\n\nPhase IV: Feedback & Closure").bold = True
    flow.add_run("\nSystem updates Timeline -> Citizen tracks via Token -> Admin reviews resolution proof -> Closes complaint.")

    # --- 9. MODULES ---
    doc.add_heading('9. Module Specifications', level=1)
    doc.add_paragraph("The system is divided into four primary functional modules:")
    
    mod = doc.add_paragraph()
    mod.add_run("1. Public Module: ").bold = True
    mod.add_run("Citizen interaction point (Submission, Mobile-responsive UI, Real-time Tracker).\n")
    mod.add_run("2. Administration Module: ").bold = True
    mod.add_run("Command center for management (Master Data, Staff Management, Global Analytics).\n")
    mod.add_run("3. Field Staff Module: ").bold = True
    mod.add_run("Operational dashboard for issue resolution and evidence management.\n")
    mod.add_run("4. Security Module: ").bold = True
    mod.add_run("Supabase Auth integration, RBAC enforcement, and secure Session management.")

    # --- 10. SYSTEM REQUIREMENTS ---
    doc.add_heading('10. Hardware & Software Requirements', level=1)
    req_table = doc.add_table(rows=4, cols=2)
    req_data = [
        ["Hardware Parameter", "Specification"],
        ["Processor", "Intel Core i3 10th Gen or higher"],
        ["Memory (RAM)", "4GB Minimum / 8GB Recommended"],
        ["Browser Support", "Chrome 90+, Firefox 88+, Safari 14+"]
    ]
    for i, row in enumerate(req_data):
        req_table.cell(i, 0).text = row[0]
        req_table.cell(i, 1).text = row[1]
        req_table.cell(i, 0).paragraphs[0].runs[0].font.bold = True

    # --- 11. OUTPUT SCREENS & CONCLUSION ---
    doc.add_heading('11. System Interfaces', level=1)
    doc.add_paragraph("The application features a modern, intuitive UI designed for high utility:")
    doc.add_paragraph("• Dashboard: Real-time status visualization.\n• Tracker: Interactive timeline of complaint history.\n• Admin Panel: Centralized management of complex state transitions.")

    doc.add_heading('12. Conclusion', level=1)
    doc.add_paragraph(
        "The Public Complaint Management System represents a significant step towards "
        "e-Governance. By automating the lifecycle of public grievances, the system "
        "ensures accountability and improves the quality of service delivery. Its scalable "
        "architecture ensures it can grow with the needs of the city."
    )

    # --- FINAL SAVING ---
    file_path = "Public_Complaint_Management_System_Professional_Doc.docx"
    doc.save(file_path)
    print(f"Documentation generated successfully: {os.path.abspath(file_path)}")

if __name__ == "__main__":
    create_professional_doc()